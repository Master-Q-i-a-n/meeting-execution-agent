from dataclasses import dataclass, field
from io import BytesIO
from pathlib import PurePath
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.llm.dashscope import TranscribedAudioSegment, ocr_meeting_image, transcribe_audio_file

MAX_TEXT_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_IMAGE_UPLOAD_BYTES = 20 * 1024 * 1024
MAX_AUDIO_UPLOAD_BYTES = 7 * 1024 * 1024
MAX_MEETING_UPLOAD_BYTES = MAX_AUDIO_UPLOAD_BYTES
SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
SUPPORTED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
SUPPORTED_AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a"}
SUPPORTED_MEETING_EXTENSIONS = (
    SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_IMAGE_EXTENSIONS | SUPPORTED_AUDIO_EXTENSIONS
)
IMAGE_MIME_TYPES = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
}
AUDIO_MIME_TYPES = {
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".m4a": "audio/mp4",
}


class MeetingContentError(ValueError):
    """会议内容解析失败时抛出的业务异常。"""


class UnsupportedMeetingFileType(MeetingContentError):
    """上传文件类型不在当前阶段支持范围内。"""


class EmptyMeetingContent(MeetingContentError):
    """解析后没有得到有效会议文本。"""


@dataclass(frozen=True)
class ParsedMeetingContent:
    """不同输入格式统一后的会议内容。

    后续 LLM 解析、分块入 Qdrant、追问检索都优先消费这里的 text。
    """

    source_type: str
    text: str
    metadata: dict[str, Any]
    audio_segments: list["ParsedAudioSegment"] = field(default_factory=list)


@dataclass(frozen=True)
class ParsedAudioSegment:
    """ASR 片段在入库前的标准形态。"""

    text: str
    start_time: float | None
    end_time: float | None
    speaker: str | None
    emotion: str | None
    pause_before_ms: int | None
    speech_rate: str | None
    confidence: float | None
    source_filename: str | None
    order_index: int


def normalize_pasted_content(content: str) -> str:
    """清洗用户直接粘贴的会议纪要文本。"""
    normalized = _normalize_text(content)
    if not normalized:
        raise EmptyMeetingContent("meeting content is empty")
    return normalized


def parse_uploaded_meeting_file(
    *,
    filename: str,
    content: bytes,
    content_type: str | None = None,
) -> ParsedMeetingContent:
    """把上传的 txt/Markdown/PDF/Word 文件解析成纯文本。

    阶段二先只做“格式统一”，暂时不在这里做摘要、待办抽取。
    """
    if not content:
        raise EmptyMeetingContent("uploaded file is empty")

    suffix = PurePath(filename).suffix.lower()
    if suffix not in SUPPORTED_MEETING_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_MEETING_EXTENSIONS))
        raise UnsupportedMeetingFileType(f"unsupported meeting file type: {suffix or 'unknown'}; supported: {supported}")

    audio_segments: list[ParsedAudioSegment] = []
    # 先按文件类型抽取文本，后面的会议分析只面对统一的纯文本。
    if suffix == ".txt":
        source_type = "text"
        text = _decode_text_bytes(content)
    elif suffix in {".md", ".markdown"}:
        source_type = "markdown"
        text = _decode_text_bytes(content)
    elif suffix == ".pdf":
        source_type = "pdf"
        text = _extract_pdf_text(content)
    elif suffix == ".docx":
        # python-docx 只支持 docx；老式 .doc 后续可通过转换服务补上。
        source_type = "word"
        text = _extract_docx_text(content)
    elif suffix in SUPPORTED_IMAGE_EXTENSIONS:
        source_type = "image"
        text = ocr_meeting_image(
            content=content,
            mime_type=_resolve_mime_type(suffix=suffix, content_type=content_type),
        )
    else:
        source_type = "audio"
        audio_segments = _enrich_audio_segments(
            transcribe_audio_file(
                content=content,
                mime_type=_resolve_mime_type(suffix=suffix, content_type=content_type),
            ),
            source_filename=filename,
        )
        text = build_audio_meeting_text(audio_segments)

    # 不同解析器抽出的换行和空白风格不一致，入库前统一清洗一遍。
    normalized = _normalize_text(text)
    if not normalized:
        raise EmptyMeetingContent("no readable meeting content was extracted")

    # metadata 留下输入来源，后续排查“这段正文从哪里来”会轻松很多。
    return ParsedMeetingContent(
        source_type=source_type,
        text=normalized,
        metadata={
            "input_mode": "upload",
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(content),
            "parser": _metadata_parser_for_source_type(source_type),
            "audio_segment_count": len(audio_segments) if source_type == "audio" else 0,
        },
        audio_segments=audio_segments if source_type == "audio" else [],
    )


def get_upload_size_limit(filename: str) -> int:
    """按文件类型返回上传大小限制，音频通常比文本纪要大很多。"""
    suffix = PurePath(filename).suffix.lower()
    if suffix in SUPPORTED_AUDIO_EXTENSIONS:
        return MAX_AUDIO_UPLOAD_BYTES
    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        return MAX_IMAGE_UPLOAD_BYTES
    return MAX_TEXT_UPLOAD_BYTES


def build_paste_metadata(content: str) -> dict[str, Any]:
    """构造粘贴文本的 metadata，方便后续追踪来源。"""
    return {
        "input_mode": "paste",
        "size_chars": len(content),
    }


def _decode_text_bytes(content: bytes) -> str:
    """尽量兼容常见中文会议纪要编码。"""
    # utf-8-sig 先处理带 BOM 的文本，gb18030 兼容常见中文本地文件。
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise MeetingContentError("text file encoding is not supported")


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise MeetingContentError(f"failed to read PDF file: {exc}") from exc

    # pypdf 按页抽取文本；扫描版 PDF 需要 OCR，放到后续阶段。
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page_text for page_text in pages if page_text.strip())


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise MeetingContentError(f"failed to read Word docx file: {exc}") from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    # 表格里也经常有负责人/截止时间，所以把单元格文本一起保留下来。
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def build_audio_meeting_text(segments: list[ParsedAudioSegment]) -> str:
    """把语音片段拼成现有会议分析流程可消费的文本。"""
    lines = []
    for segment in segments:
        time_range = _format_time_range(segment.start_time, segment.end_time)
        pause = f"[pause={segment.pause_before_ms}ms]" if segment.pause_before_ms else ""
        emotion = f"[{segment.emotion}]" if segment.emotion else ""
        speech_rate = f"[{segment.speech_rate}]" if segment.speech_rate else ""
        speaker = f"{segment.speaker}：" if segment.speaker else ""
        lines.append(f"{time_range}{pause}{emotion}{speech_rate} {speaker}{segment.text}".strip())
    return "\n".join(lines)


def _enrich_audio_segments(
    segments: list[TranscribedAudioSegment],
    *,
    source_filename: str,
) -> list[ParsedAudioSegment]:
    enriched: list[ParsedAudioSegment] = []
    previous_end: float | None = None
    for index, segment in enumerate(segments):
        pause_before_ms = _calculate_pause_before_ms(previous_end, segment.start_time)
        speech_rate = _calculate_speech_rate_label(
            text=segment.text,
            start_time=segment.start_time,
            end_time=segment.end_time,
        )
        enriched.append(
            ParsedAudioSegment(
                text=segment.text,
                start_time=segment.start_time,
                end_time=segment.end_time,
                speaker=segment.speaker,
                emotion=segment.emotion,
                pause_before_ms=pause_before_ms,
                speech_rate=speech_rate,
                confidence=segment.confidence,
                source_filename=source_filename,
                order_index=index,
            )
        )
        if segment.end_time is not None:
            previous_end = segment.end_time
    return enriched


def _calculate_pause_before_ms(previous_end: float | None, current_start: float | None) -> int | None:
    if previous_end is None or current_start is None:
        return None
    pause_ms = int(max(current_start - previous_end, 0) * 1000)
    return pause_ms if pause_ms > 0 else 0


def _calculate_speech_rate_label(
    *,
    text: str,
    start_time: float | None,
    end_time: float | None,
) -> str | None:
    if start_time is None or end_time is None or end_time <= start_time:
        return None
    chars_per_second = len(text.strip()) / max(end_time - start_time, 0.1)
    if chars_per_second < 2:
        return "slow"
    if chars_per_second > 5:
        return "fast"
    return "normal"


def _format_time_range(start_time: float | None, end_time: float | None) -> str:
    if start_time is None and end_time is None:
        return "[unknown-time]"
    return f"[{_format_seconds(start_time)}-{_format_seconds(end_time)}]"


def _format_seconds(value: float | None) -> str:
    if value is None:
        return "??:??"
    total_seconds = int(value)
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def _resolve_mime_type(*, suffix: str, content_type: str | None) -> str:
    # 多模态模型对 Data URL 的 MIME 比较敏感，优先用后缀映射出的标准类型。
    return (
        IMAGE_MIME_TYPES.get(suffix)
        or AUDIO_MIME_TYPES.get(suffix)
        or (content_type if content_type and content_type != "application/octet-stream" else None)
        or "application/octet-stream"
    )


def _metadata_parser_for_source_type(source_type: str) -> str:
    if source_type == "image":
        return "dashscope_vision_ocr"
    if source_type == "audio":
        return "dashscope_asr"
    return "local_text_extractor"


def _normalize_text(content: str) -> str:
    # 数据库中统一保存 \n，避免 Windows 和其他来源的换行符混杂。
    lines = [line.rstrip() for line in content.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(lines).strip()
