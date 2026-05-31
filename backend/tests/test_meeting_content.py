from io import BytesIO

import pytest
from docx import Document

from app.services.meeting_content import (
    MAX_AUDIO_UPLOAD_BYTES,
    MAX_IMAGE_UPLOAD_BYTES,
    MAX_TEXT_UPLOAD_BYTES,
    EmptyMeetingContent,
    ParsedAudioSegment,
    UnsupportedMeetingFileType,
    build_audio_meeting_text,
    get_upload_size_limit,
    normalize_pasted_content,
    parse_uploaded_meeting_file,
)


def test_normalize_pasted_content_strips_outer_blank_lines() -> None:
    """粘贴文本会被清理成后续解析更稳定的格式。"""
    content = normalize_pasted_content("\n\n会议结论：继续推进\r\n负责人：张三  \n")

    assert content == "会议结论：继续推进\n负责人：张三"


def test_parse_markdown_file_as_text() -> None:
    """Markdown 阶段二先作为纯文本保存，后续再做结构化解析。"""
    parsed = parse_uploaded_meeting_file(
        filename="weekly.md",
        content="# 周会\n\n- 李四负责接口联调".encode(),
        content_type="text/markdown",
    )

    assert parsed.source_type == "markdown"
    assert "李四负责接口联调" in parsed.text
    assert parsed.metadata["filename"] == "weekly.md"


def test_parse_docx_file_extracts_paragraph_text() -> None:
    """Word docx 可以抽取段落文本，先覆盖最常见会议纪要格式。"""
    document = Document()
    document.add_paragraph("项目复盘会议")
    document.add_paragraph("王五下周三前完成风险清单")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_uploaded_meeting_file(
        filename="review.docx",
        content=buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.source_type == "word"
    assert "王五下周三前完成风险清单" in parsed.text


def test_parse_image_file_uses_ocr(monkeypatch: pytest.MonkeyPatch) -> None:
    """图片会议先走 OCR，再作为统一会议文本进入后续流程。"""
    monkeypatch.setattr(
        "app.services.meeting_content.ocr_meeting_image",
        lambda **kwargs: "手写纪要：张三负责接口联调。",
    )

    parsed = parse_uploaded_meeting_file(
        filename="notes.png",
        content=b"fake-image",
        content_type="image/png",
    )

    assert parsed.source_type == "image"
    assert parsed.text == "手写纪要：张三负责接口联调。"
    assert parsed.metadata["parser"] == "dashscope_vision_ocr"


def test_parse_audio_file_generates_segments(monkeypatch: pytest.MonkeyPatch) -> None:
    """音频会议会保存 ASR 片段，并拼成带时间戳的会议原文。"""
    monkeypatch.setattr(
        "app.services.meeting_content.transcribe_audio_file",
        lambda **kwargs: [
            ParsedAudioSegment(
                text="张三负责接口联调。",
                start_time=12.0,
                end_time=18.0,
                speaker="speaker_1",
                emotion="neutral",
                pause_before_ms=None,
                speech_rate=None,
                confidence=0.9,
                source_filename="meeting.mp3",
                order_index=0,
            ),
            ParsedAudioSegment(
                text="下周三前完成。",
                start_time=19.0,
                end_time=21.0,
                speaker="speaker_1",
                emotion="neutral",
                pause_before_ms=None,
                speech_rate=None,
                confidence=0.9,
                source_filename="meeting.mp3",
                order_index=1,
            ),
        ],
    )

    parsed = parse_uploaded_meeting_file(
        filename="meeting.mp3",
        content=b"fake-audio",
        content_type="audio/mpeg",
    )

    assert parsed.source_type == "audio"
    assert parsed.metadata["audio_segment_count"] == 2
    assert parsed.audio_segments[1].pause_before_ms == 1000
    assert "[00:00:12-00:00:18]" in parsed.text
    assert "下周三前完成" in parsed.text


def test_build_audio_meeting_text_includes_voice_clues() -> None:
    text = build_audio_meeting_text(
        [
            ParsedAudioSegment(
                text="这个需求可能做不完。",
                start_time=1.0,
                end_time=5.0,
                speaker=None,
                emotion="neutral",
                pause_before_ms=1800,
                speech_rate="slow",
                confidence=0.8,
                source_filename="meeting.mp3",
                order_index=0,
            )
        ]
    )

    assert "[pause=1800ms]" in text
    assert "[neutral]" in text
    assert "[slow]" in text


def test_parse_uploaded_file_rejects_unsupported_extension() -> None:
    """当前阶段不支持的文件类型要明确报错。"""
    with pytest.raises(UnsupportedMeetingFileType):
        parse_uploaded_meeting_file(filename="archive.zip", content=b"fake")


def test_parse_uploaded_file_rejects_empty_text() -> None:
    """空文件或解析后为空的文件不能进入后续 Agent 流程。"""
    with pytest.raises(EmptyMeetingContent):
        parse_uploaded_meeting_file(filename="empty.txt", content=b"")


def test_upload_size_limit_depends_on_file_type() -> None:
    """音频文件允许比普通文档更大的上传体积。"""
    assert get_upload_size_limit("meeting.md") == MAX_TEXT_UPLOAD_BYTES
    assert get_upload_size_limit("whiteboard.png") == MAX_IMAGE_UPLOAD_BYTES
    assert get_upload_size_limit("recording.mp3") == MAX_AUDIO_UPLOAD_BYTES
